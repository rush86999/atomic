import io
import os
import logging
from typing import List, Dict, Any, Optional
from pdfminer.high_level import extract_text_to_fp
from pdfminer.layout import LAParams
from datetime import datetime, timezone # For ingested_at timestamp

# Attempt to import python-docx
try:
    import docx # from python-docx
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    logging.getLogger(__name__).warning("python-docx library not found. DOCX processing will not be available.")
    PYTHON_DOCX_AVAILABLE = False
    docx = None

# Embedding and LanceDB imports (adjust paths as necessary based on execution context)
try:
    import sys
    FUNCTIONS_DIR_PATH = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', 'project', 'functions'))
    if FUNCTIONS_DIR_PATH not in sys.path:
        sys.path.append(FUNCTIONS_DIR_PATH)
    from note_utils import get_text_embedding_openai
    EMBEDDING_FUNCTION_AVAILABLE = True
except ImportError as e:
    logging.getLogger(__name__).warning(f"Could not import get_text_embedding_openai from note_utils: {e}. Embedding will not be available.")
    get_text_embedding_openai = None
    EMBEDDING_FUNCTION_AVAILABLE = False

try:
    from . import lancedb_handler
except ImportError:
    import lancedb_handler # Fallback for different execution contexts
    logging.getLogger(__name__).warning("Imported lancedb_handler using fallback.")


logger = logging.getLogger(__name__)
if not logger.hasHandlers():
    logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')

# --- Text Extraction Utilities ---
def extract_text_from_pdf(pdf_file_path_or_bytes: Any) -> str:
    try:
        output_string = io.StringIO()
        laparams = LAParams()
        if isinstance(pdf_file_path_or_bytes, str):
            if not os.path.exists(pdf_file_path_or_bytes):
                raise FileNotFoundError(f"PDF file not found at path: {pdf_file_path_or_bytes}")
            with open(pdf_file_path_or_bytes, 'rb') as fin:
                extract_text_to_fp(fin, output_string, laparams=laparams, output_type='text', codec='utf-8')
        elif isinstance(pdf_file_path_or_bytes, bytes):
             with io.BytesIO(pdf_file_path_or_bytes) as fin:
                extract_text_to_fp(fin, output_string, laparams=laparams, output_type='text', codec='utf-8')
        else:
            raise ValueError("Input for PDF extraction must be a file path (str) or bytes.")
        return output_string.getvalue()
    except Exception as e:
        logger.error(f"Error extracting text from PDF input: {e}", exc_info=True)
        raise

def extract_text_from_docx(file_path_or_bytes: Any) -> str:
    if not PYTHON_DOCX_AVAILABLE or not docx:
        raise ImportError("python-docx library is required for DOCX processing but not found.")
    try:
        if isinstance(file_path_or_bytes, str):
            if not os.path.exists(file_path_or_bytes):
                raise FileNotFoundError(f"DOCX file not found at path: {file_path_or_bytes}")
            document = docx.Document(file_path_or_bytes)
        elif isinstance(file_path_or_bytes, bytes):
            document = docx.Document(io.BytesIO(file_path_or_bytes))
        else:
            raise ValueError("Input for DOCX extraction must be a file path (str) or bytes.")
        full_text = [para.text for para in document.paragraphs if para.text.strip()]
        return '\n\n'.join(full_text) # Join paragraphs with double newline for some separation
    except Exception as e:
        logger.error(f"Error extracting text from DOCX input: {e}", exc_info=True)
        raise

def extract_text_from_txt(file_path_or_bytes: Any, encoding: str = 'utf-8') -> str:
    try:
        if isinstance(file_path_or_bytes, str):
            if not os.path.exists(file_path_or_bytes):
                raise FileNotFoundError(f"TXT file not found at path: {file_path_or_bytes}")
            with open(file_path_or_bytes, 'r', encoding=encoding) as f:
                return f.read()
        elif isinstance(file_path_or_bytes, bytes):
            return file_path_or_bytes.decode(encoding)
        else:
            raise ValueError("Input for TXT extraction must be a file path (str) or bytes.")
    except UnicodeDecodeError as ude:
        logger.error(f"Unicode error for TXT (encoding: {encoding}): {ude}", exc_info=True)
        raise ValueError(f"Encoding error: Could not decode TXT with {encoding}.") from ude
    except Exception as e:
        logger.error(f"Error extracting text from TXT input: {e}", exc_info=True)
        raise

# --- Text Chunking Utility ---
def chunk_text(text: str, chunk_size: int = 1000, chunk_overlap: int = 100) -> List[str]:
    if not text: return []
    chunks = []; start = 0; text_len = len(text)
    while start < text_len:
        end = start + chunk_size
        chunks.append(text[start:end])
        if end >= text_len: break
        start += (chunk_size - chunk_overlap)
        if start >= text_len: break
    return chunks

# --- Main Orchestration Function ---
async def process_document_and_store(
    user_id: str,
    file_path_or_bytes: Any, # Can be path (str) or content (bytes)
    document_id: str,
    source_uri: str, # Original URI of the document
    original_doc_type: str, # Original type, e.g., "gdoc", "pdf_upload", "gdrive_docx"
    processing_mime_type: str, # MIME type of file_path_or_bytes (e.g., "application/pdf", "text/plain")
    title: Optional[str] = None,
    doc_metadata_json: Optional[str] = None,
    openai_api_key_param: Optional[str] = None
) -> Dict[str, Any]:
    """
    Orchestrates document processing: text extraction based on MIME type, chunking,
    embedding, and storage in LanceDB.
    """
    logger.info(f"Processing document: id={document_id}, user={user_id}, source_uri={source_uri}, original_type='{original_doc_type}', processing_mime='{processing_mime_type}'")

    if not EMBEDDING_FUNCTION_AVAILABLE or not get_text_embedding_openai:
        return {"status": "error", "message": "Embedding utility (note_utils/get_text_embedding_openai) not available.", "code": "EMBEDDING_SERVICE_UNAVAILABLE"}
    if not lancedb_handler:
         return {"status": "error", "message": "LanceDB handler not available.", "code": "LANCEDB_HANDLER_UNAVAILABLE"}

    extracted_text = ""
    try:
        if processing_mime_type == "application/pdf":
            extracted_text = extract_text_from_pdf(file_path_or_bytes)
        elif processing_mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or \
             original_doc_type.endswith("docx"): # Broader check for docx
            extracted_text = extract_text_from_docx(file_path_or_bytes)
        elif processing_mime_type == "text/plain" or original_doc_type.endswith("txt"):
            extracted_text = extract_text_from_txt(file_path_or_bytes)
        else:
            logger.error(f"Unsupported processing_mime_type: {processing_mime_type} for document {document_id}")
            # Update document metadata in LanceDB to 'failed_unsupported_type'
            # This requires doc_meta to be created first, then updated.
            # For now, just return error. A more robust flow would pre-register the doc then update status.
            return {"status": "error", "message": f"Unsupported document type for processing: {processing_mime_type}", "code": "UNSUPPORTED_PROCESSING_TYPE"}
    except Exception as extraction_error:
        logger.error(f"Text extraction failed for document {document_id} (type: {processing_mime_type}): {extraction_error}", exc_info=True)
        # TODO: Update document status in LanceDB to "failed_extraction"
        return {"status": "error", "message": f"Text extraction failed: {str(extraction_error)}", "code": "TEXT_EXTRACTION_FAILED"}

    if not extracted_text.strip():
        logger.warning(f"No text extracted from document: {source_uri} (ID: {document_id})")
        doc_meta_to_store = {
            "doc_id": document_id, "user_id": user_id, "source_uri": source_uri,
            "doc_type": original_doc_type, "title": title, "metadata_json": doc_metadata_json,
            "ingested_at": datetime.now(timezone.utc).isoformat(), # Store as ISO string
            "processing_status": "failed_empty_content", "error_message": "No text content extracted."
        }
        db_conn = await lancedb_handler.get_lancedb_connection()
        if db_conn:
            await lancedb_handler.create_generic_document_tables_if_not_exist(db_conn)
            await lancedb_handler.add_processed_document(db_conn, doc_meta_to_store, []) # Store metadata even if no chunks
        return {"status": "warning", "message": "No text extracted from document, only metadata stored.", "code": "PDF_PROCESSING_EMPTY", "data": {"doc_id": document_id, "num_chunks_stored": 0}}

    text_chunks = chunk_text(extracted_text)
    logger.info(f"Extracted {len(extracted_text)} chars, split into {len(text_chunks)} chunks for doc_id: {document_id}")

    chunks_with_embeddings_data = []
    for i, chunk_text_content in enumerate(text_chunks):
        if not chunk_text_content.strip(): continue
        embedding_resp = get_text_embedding_openai(text_to_embed=chunk_text_content, openai_api_key_param=openai_api_key_param)
        if embedding_resp["status"] == "success":
            chunks_with_embeddings_data.append({
                "doc_id": document_id, # Ensure doc_id is part of the chunk data
                "user_id": user_id,   # Ensure user_id is part of the chunk data
                "parent_doc_type": original_doc_type, # Add parent_doc_type
                "chunk_sequence": i,
                "text_content": chunk_text_content,
                "embedding": embedding_resp["data"], # Ensure key is 'embedding'
                "metadata_json": None,
                "char_count": len(chunk_text_content)
            })
        else:
            logger.warning(f"Embedding failed for chunk {i} of doc_id {document_id}: {embedding_resp.get('message')}")

    if not chunks_with_embeddings_data:
        doc_meta_err_embed = {
            "doc_id": document_id, "user_id": user_id, "source_uri": source_uri,
            "doc_type": original_doc_type, "title": title, "metadata_json": doc_metadata_json,
            "ingested_at": datetime.now(timezone.utc).isoformat(),
            "processing_status": "failed_embedding", "error_message": "No chunks could be embedded."
        }
        db_conn = await lancedb_handler.get_lancedb_connection()
        if db_conn:
            await lancedb_handler.create_generic_document_tables_if_not_exist(db_conn)
            await lancedb_handler.add_processed_document(db_conn, doc_meta_err_embed, [])
        return {"status": "error", "message": "No chunks could be embedded for the document.", "code": "EMBEDDING_FAILURE_ALL_CHUNKS"}

    doc_meta_to_store = {
        "doc_id": document_id, "user_id": user_id, "source_uri": source_uri,
        "doc_type": original_doc_type, "title": title, "metadata_json": doc_metadata_json,
        "ingested_at": datetime.now(timezone.utc).isoformat(),
        "processing_status": "completed",
    }

    db_conn = await lancedb_handler.get_lancedb_connection()
    if not db_conn:
        return {"status": "error", "message": "Failed to connect to LanceDB.", "code": "LANCEDB_CONNECTION_ERROR"}

    await lancedb_handler.create_generic_document_tables_if_not_exist(db_conn)
    storage_result = await lancedb_handler.add_processed_document(
        db_conn=db_conn, doc_meta=doc_meta_to_store, chunks_with_embeddings=chunks_with_embeddings_data
    )

    if storage_result["status"] == "success":
        logger.info(f"Successfully processed and stored document: {source_uri}, doc_id: {document_id}")
        return {"status": "success", "data": {"doc_id": document_id, "num_chunks_stored": len(chunks_with_embeddings_data)}}
    else:
        logger.error(f"Failed to store processed document {document_id} in LanceDB: {storage_result.get('message')}")
        return {"status": "error", "message": f"LanceDB storage failed: {storage_result.get('message')}", "code": storage_result.get("code", "LANCEDB_STORAGE_ERROR")}

```
